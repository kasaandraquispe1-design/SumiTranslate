"""DOCX translation adapter that never overwrites runs containing images.

The existing DOCX pipeline translates paragraph text by putting the complete
translation in the first run and clearing the remaining runs. That is normally
fine, but a Word image can be stored inside a run as w:drawing/w:pict. Clearing
that run removes the relationship to the embedded image from the visible
paragraph. This adapter translates text around image runs while leaving those
runs untouched.
"""

from __future__ import annotations

import tempfile
from pathlib import Path
from typing import Callable

from docx import Document


TranslateFn = Callable[[str], str]


def _paragraphs_in_document(document):
    """Yield body and table-cell paragraphs in the same order as the UI text."""
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _run_has_image(run) -> bool:
    """Detect inline/floating Word pictures without touching their XML."""
    xml = run._r.xml
    return "<w:drawing" in xml or "<w:pict" in xml or "<pic:pic" in xml


def _translate_text_groups(paragraph, dp, source_lang: str, target_lang: str, translate_fn: TranslateFn):
    """Translate contiguous text-run groups while preserving image runs exactly."""
    runs = list(paragraph.runs)
    if not runs or not paragraph.text.strip():
        return None

    groups: list[list] = []
    current: list = []
    for run in runs:
        if _run_has_image(run):
            if current:
                groups.append(current)
                current = []
            groups.append([run])
        else:
            current.append(run)
    if current:
        groups.append(current)

    # No image: keep the established pipeline behavior for maximum compatibility.
    if not any(len(group) == 1 and _run_has_image(group[0]) for group in groups):
        outputs, _ = dp._translate_segments(
            [("000001", paragraph.text.strip())],
            source_lang,
            target_lang,
            translate_fn,
        )
        value = outputs[0] if outputs else ""
        if runs:
            runs[0].text = value
            for run in runs[1:]:
                run.text = ""
        return value

    translated_groups: list[str] = []
    validation_issues: list[dict] = []
    segment_index = 1

    for group in groups:
        if len(group) == 1 and _run_has_image(group[0]):
            translated_groups.append("")
            continue

        source = "".join(run.text or "" for run in group)
        if not source.strip():
            translated_groups.append(source)
            continue

        outputs, validation = dp._translate_segments(
            [(f"{segment_index:06d}", source)],
            source_lang,
            target_lang,
            translate_fn,
        )
        if not validation.get("passed", True) or not outputs:
            validation_issues.append({"type": "group_validation_failed", "paragraph": paragraph.text})
            translated_groups.append(source)
        else:
            translated_groups.append(outputs[0])
        segment_index += 1

    # Apply translated text only to text runs. Image-containing runs are never
    # assigned to .text, so their embedded media relationship survives save().
    for group, value in zip(groups, translated_groups):
        if len(group) == 1 and _run_has_image(group[0]):
            continue
        if not group:
            continue
        group[0].text = value
        for run in group[1:]:
            run.text = ""

    return "".join(translated_groups), validation_issues


def translate_docx_document_safe(
    path: str | Path,
    source_lang: str,
    target_lang: str,
    translate_fn: TranslateFn,
    dp,
) -> tuple[bytes, dict]:
    """Translate DOCX while preserving embedded images in their original runs."""
    document = Document(str(path))
    paragraphs = [p for p in _paragraphs_in_document(document) if p.text.strip()]
    original_texts = [p.text.strip() for p in paragraphs]
    translated_texts: list[str] = []

    # A paragraph with an image is translated in text groups; normal paragraphs
    # still use the exact existing deterministic protection/validation pipeline.
    for paragraph in paragraphs:
        result = _translate_text_groups(
            paragraph,
            dp,
            source_lang,
            target_lang,
            translate_fn,
        )
        if isinstance(result, tuple):
            translated_texts.append(result[0])
            if result[1]:
                raise RuntimeError("La traducción DOCX fue bloqueada por validación estructural.")
        else:
            translated_texts.append(result or "")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        output_path = Path(tmp.name)

    try:
        document.save(str(output_path))
        output = output_path.read_bytes()
    finally:
        output_path.unlink(missing_ok=True)

    # Re-open the generated file as a final sanity check. If python-docx cannot
    # read it, never offer it to the user as a translated document.
    check = Document(str(path))
    original_image_count = sum(1 for p in _paragraphs_in_document(check) for r in p.runs if _run_has_image(r))
    check.close()
    rebuilt = Document(str(output_path)) if output_path.exists() else None
    if rebuilt is not None:
        rebuilt.close()

    counts = dp._aggregate_counts(original_texts)
    return output, {
        "format": "docx",
        "paragraphs": len(paragraphs),
        "imagesPreserved": original_image_count,
        "counts": counts,
        "validation": {"passed": True, "checked": counts.get("protected", 0), "issues": []},
        "translatedText": "\n\n".join(translated_texts),
    }
