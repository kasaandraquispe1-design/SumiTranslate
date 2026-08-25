"""Layout-aware document translation for Sumire Translate."""

from __future__ import annotations

import re
import zipfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Callable

from backend.core.languages import language_name
from backend.protection.math_protector import ProtectedElement, protect_text, restore_markers
from backend.processing.word_counter import count_words
from backend.validation.validator import validate

TranslateFn = Callable[[str], str]
MARKER_RE = re.compile(r"\[\[[A-Z]+_\d{3}\]\]")

@dataclass
class PDFTextBlock:
    page: int
    bbox: tuple[float, float, float, float]
    text: str
    fontsize: float
    flags: int
    color: int

def _block_text(block: dict) -> str:
    lines = []
    for line in block.get("lines", []):
        spans = line.get("spans", [])
        lines.append("".join(str(span.get("text", "")) for span in spans).rstrip())
    return "\n".join(lines).strip()

def _extract_pdf_blocks(doc):
    blocks = []
    for page_index, page in enumerate(doc):
        data = page.get_text("dict", sort=True)
        for block in data.get("blocks", []):
            if block.get("type") != 0 or not block.get("lines"):
                continue
            text = _block_text(block)
            if not text:
                continue
            spans = [span for line in block.get("lines", []) for span in line.get("spans", []) if span.get("text")]
            if not spans:
                continue
            first = spans[0]
            blocks.append(PDFTextBlock(page_index, tuple(float(x) for x in block["bbox"]), text,
                                       float(first.get("size", 10.0) or 10.0),
                                       int(first.get("flags", 0) or 0), int(first.get("color", 0) or 0)))
    return blocks

def _table_cell_rects(table):
    rects = []
    for cell in getattr(table, "cells", []) or []:
        try:
            if cell is not None and len(cell) >= 4:
                rects.append(tuple(float(v) for v in cell[:4]))
        except (TypeError, ValueError):
            continue
    return rects

def _infer_matrix(rects):
    if not rects:
        return 0, 0
    ys = sorted({round((r[1] + r[3]) / 2, 2) for r in rects})
    xs = sorted({round((r[0] + r[2]) / 2, 2) for r in rects})
    return len(ys), len(xs)

def _detect_pdf_tables(doc):
    tables = []
    for page_index, page in enumerate(doc):
        try:
            found = getattr(page.find_tables(), "tables", []) or []
        except (AttributeError, TypeError, RuntimeError):
            found = []
        for table_index, table in enumerate(found, start=1):
            try:
                extracted = [list(row) for row in (table.extract() or [])]
            except (AttributeError, TypeError, RuntimeError):
                extracted = []
            rows = len(extracted)
            columns = max((len(row) for row in extracted), default=0)
            rects = _table_cell_rects(table)
            if not rows or not columns:
                rows, columns = _infer_matrix(rects)
            bbox = getattr(table, "bbox", None)
            tables.append({"number": table_index, "page": page_index + 1, "rows": rows,
                           "columns": columns, "cells": rows * columns if rows and columns else len(rects),
                           "size": f"{rows}×{columns}" if rows and columns else "desconocido",
                           "bbox": tuple(float(x) for x in bbox) if bbox else None})
    return tables

def inspect_pdf_tables(path):
    try:
        import pymupdf
    except ImportError as exc:
        raise RuntimeError("Falta PyMuPDF. La dependencia está declarada en requirements.txt.") from exc
    doc = pymupdf.open(str(path))
    try:
        return _detect_pdf_tables(doc)
    finally:
        doc.close()

def _marker_sequence(text):
    return MARKER_RE.findall(text or "")

def _build_segment_prompt(source_lang, target_lang, segments):
    body = []
    for segment_id, protected in segments:
        body.extend([f"[[SUMIRE_SEG_{segment_id}_START]]", protected, f"[[SUMIRE_SEG_{segment_id}_END]]"])
    return f"""You are a professional academic and scientific translator.
Translate from {language_name(source_lang)} to {language_name(target_lang)}.

Translate ONLY natural language. Preserve every marker exactly.
Never translate, modify, reorder, remove or add protected markers.
Do not change formulas, variables, numbers, symbols, code, URLs, citations,
table delimiters or document structure.
Keep every segment delimiter exactly as supplied.
Output only the translation.

{chr(10).join(body)}"""

def _normalise_segments(segments):
    normalised = []
    for index, item in enumerate(segments, start=1):
        if isinstance(item, (tuple, list)) and len(item) == 2:
            segment_id, text = item
        elif isinstance(item, str):
            segment_id, text = f"{index:06d}", item
        else:
            raise RuntimeError(f"Segmento DOCX inválido en posición {index}: se esperaba (id, texto).")
        normalised.append((str(segment_id), str(text)))
    return normalised

def _translate_segments(segments, source_lang, target_lang, translate_fn):
    segments = _normalise_segments(segments)
    if not segments:
        return []
    prepared = []
    for segment_id, text in segments:
        protected, store, _ = protect_text(text)
        prepared.append((segment_id, protected, store))

    translated = translate_fn(_build_segment_prompt(
        source_lang, target_lang, [(segment_id, protected) for segment_id, protected, _ in prepared]
    ))
    if not isinstance(translated, str):
        translated = str(translated)

    expected = []
    for segment_id, protected, _ in prepared:
        expected.append(f"[[SUMIRE_SEG_{segment_id}_START]]")
        expected.extend(_marker_sequence(protected))
        expected.append(f"[[SUMIRE_SEG_{segment_id}_END]]")
    actual = re.findall(r"\[\[[A-Z_]+_\d{3,6}(?:_(?:START|END))?\]\]", translated)
    if actual != expected:
        raise RuntimeError("La traducción fue bloqueada: el modelo alteró la estructura protegida del documento.")

    outputs = []
    for segment_id, original_protected, store in prepared:
        start = f"[[SUMIRE_SEG_{segment_id}_START]]"
        end = f"[[SUMIRE_SEG_{segment_id}_END]]"
        a = translated.find(start)
        b = translated.find(end, a + len(start))
        if a < 0 or b < 0:
            raise RuntimeError("La traducción fue bloqueada: falta un delimitador de segmento.")
        segment_output = translated[a + len(start):b].strip("\n")
        if _marker_sequence(segment_output) != _marker_sequence(original_protected):
            raise RuntimeError(f"La traducción fue bloqueada: se alteraron elementos protegidos en el segmento {segment_id}.")
        restored = restore_markers(segment_output, store)
        validation = validate(original_protected.replace("[[", "").replace("]]", ""), restored, store,
                             protected_source=original_protected, translated_protected=segment_output)
        if not validation["passed"]:
            raise RuntimeError(f"La traducción fue bloqueada por la validación del segmento {segment_id}.")
        outputs.append(restored)
    return outputs

def _aggregate_counts(texts):
    total = translatable = protected = protected_words = 0
    by_type = {}
    for text in texts:
        c = count_words(text)
        total += int(c.get("total", 0)); translatable += int(c.get("translatable", 0))
        protected += int(c.get("protected", 0)); protected_words += int(c.get("protectedWords", 0))
        for kind, amount in c.get("protectedByType", {}).items():
            by_type[kind] = by_type.get(kind, 0) + int(amount)
    return {"total": total, "translatable": translatable, "protected": protected,
            "protectedWords": protected_words, "protectedRatio": protected_words / total if total else 0.0,
            "protectedByType": dict(sorted(by_type.items()))}

def _pdf_color(value):
    return ((value >> 16 & 255) / 255, (value >> 8 & 255) / 255, (value & 255) / 255)

def _pdf_font_name(flags):
    bold, italic = bool(flags & 16), bool(flags & 2)
    if bold and italic: return "figbi"
    if bold: return "figbo"
    if italic: return "figit"
    return "figo"

def _insert_pdf_text_fitted(page, rect, text, block, *, max_bottom=None):
    """Insert translated text using a controlled recovery ladder.

    Attempt 1: original rectangle + original font size.
    Attempt 2: +0.5 pt safe vertical expansion + 95% font.
    Attempt 3: +1.0 pt + 92% font.
    Attempt 4: +1.5 pt + 89% font.
    Attempt 5: +2.0 pt + 86% font.

    Expansion is capped by max_bottom so a block cannot grow into the next
    overlapping block. If an attempt fits, it is accepted immediately.
    """
    import pymupdf

    base = pymupdf.Rect(rect)
    original_bottom = base.y1
    if max_bottom is None:
        max_bottom = original_bottom
    max_bottom = max(original_bottom, float(max_bottom))

    # The requested expansion is vertical and deliberately small. The caller
    # supplies max_bottom based on the nearest overlapping block on the page.
    requested_expansions = (0.0, 0.5, 1.0, 1.5, 2.0)
    font_scales = (1.00, 0.95, 0.92, 0.89, 0.86)
    start_size = max(4.0, min(float(block.fontsize), 24.0))

    for expansion, scale in zip(requested_expansions, font_scales):
        candidate_bottom = min(original_bottom + expansion, max_bottom)
        candidate = pymupdf.Rect(base.x0, base.y0, base.x1, candidate_bottom)
        fontsize = max(4.0, start_size * scale)
        rc = page.insert_textbox(
            candidate,
            text,
            fontname=_pdf_font_name(block.flags),
            fontsize=fontsize,
            color=_pdf_color(block.color),
            overlay=True,
        )
        if rc >= 0:
            return {"fontsize": fontsize, "expansion": candidate_bottom - original_bottom}

    # A final tiny-font fallback is intentionally conservative. It does not
    # expand beyond the safe 2 pt window and is only used when the requested
    # ladder could not fit because of a borderline text measurement.
    candidate_bottom = min(original_bottom + 2.0, max_bottom)
    candidate = pymupdf.Rect(base.x0, base.y0, base.x1, candidate_bottom)
    for fontsize in (max(4.0, start_size * 0.84), max(4.0, start_size * 0.82)):
        rc = page.insert_textbox(
            candidate,
            text,
            fontname=_pdf_font_name(block.flags),
            fontsize=fontsize,
            color=_pdf_color(block.color),
            overlay=True,
        )
        if rc >= 0:
            return {"fontsize": fontsize, "expansion": candidate_bottom - original_bottom}

    raise RuntimeError("La reconstrucción fue bloqueada: un bloque traducido no cabe en su región disponible.")

def _table_cell_text(page, rect):
    try: return page.get_text("text", clip=rect, sort=True).strip()
    except (TypeError, RuntimeError): return ""

def _table_cells_for_page(page, table_obj):
    import pymupdf
    cells = []
    for rect in _table_cell_rects(table_obj):
        r = pymupdf.Rect(rect); text = _table_cell_text(page, r)
        if text: cells.append({"rect": r, "text": text})
    cells.sort(key=lambda c: (round(c["rect"].y0, 1), round(c["rect"].x0, 1)))
    return cells

def _block_inside_table(block, table_infos):
    cx = (block.bbox[0] + block.bbox[2]) / 2; cy = (block.bbox[1] + block.bbox[3]) / 2
    return any(t.get("bbox") and t["bbox"][0] <= cx <= t["bbox"][2] and t["bbox"][1] <= cy <= t["bbox"][3] and t["page"] - 1 == block.page for t in table_infos)

def _pdf_block_available_bottom(block, blocks, page):
    """Find the nearest safe lower boundary for a normal text block."""
    x0, y0, x1, y1 = block.bbox
    candidates = []
    for other in blocks:
        if other is block or other.page != block.page:
            continue
        ox0, oy0, ox1, oy1 = other.bbox
        horizontal_overlap = min(x1, ox1) - max(x0, ox0)
        if horizontal_overlap > max(2.0, min(x1 - x0, ox1 - ox0) * 0.12) and oy0 > y1:
            candidates.append(oy0 - 1.5)
    page_rect = page.rect
    candidates.append(page_rect.y1 - 3.0)
    return max(y1, min(candidates))

def _validate_final_pdf(output, source_texts):
    import pymupdf
    result = pymupdf.open(stream=output, filetype="pdf")
    extracted = "\n".join(page.get_text("text") for page in result); result.close()
    issues = []; checked = 0
    for text in source_texts:
        _, store, _ = protect_text(text)
        for item in store.values():
            checked += 1
            needle = item.original
            if needle not in extracted and re.sub(r"\s+", " ", needle).strip() not in re.sub(r"\s+", " ", extracted).strip():
                issues.append({"type": "protected_content_missing_in_pdf", "kind": item.type, "original": needle})
                if len(issues) >= 20: break
        if len(issues) >= 20: break
    return {"passed": not issues, "checked": checked, "issues": issues}

def translate_pdf_document(path, source_lang, target_lang, translate_fn, *, batch_size=18):
    try:
        import pymupdf
    except ImportError as exc:
        raise RuntimeError("Falta PyMuPDF. La dependencia está declarada en requirements.txt.") from exc
    source = pymupdf.open(str(path)); page_count = source.page_count; table_infos = _detect_pdf_tables(source)
    page_tables = {}
    for page_index, page in enumerate(source):
        try: page_tables[page_index] = list(getattr(page.find_tables(), "tables", []) or [])
        except (AttributeError, TypeError, RuntimeError): page_tables[page_index] = []
    blocks = [b for b in _extract_pdf_blocks(source) if not _block_inside_table(b, table_infos)]
    original_block_texts = [b.text for b in blocks]; original_table_texts = []; translated_blocks = []
    for start in range(0, len(blocks), batch_size):
        batch = blocks[start:start + batch_size]
        translated_blocks.extend(_translate_segments([(f"{start + i + 1:06d}", b.text) for i, b in enumerate(batch)], source_lang, target_lang, translate_fn))
    if len(translated_blocks) != len(blocks):
        source.close(); raise RuntimeError("La traducción no produjo todos los bloques del documento.")
    for page_index, table_list in page_tables.items():
        page = source[page_index]
        for table_obj in table_list:
            cells = _table_cells_for_page(page, table_obj)
            if not cells: continue
            cell_results = []
            for cell_index, cell in enumerate(cells, start=1):
                original_table_texts.append(cell["text"])
                cell_results.append((cell, _translate_segments([(f"9{page_index:02d}{cell_index:03d}", cell["text"])], source_lang, target_lang, translate_fn)[0]))
            for cell, _ in cell_results: page.add_redact_annot(cell["rect"], fill=False, cross_out=False)
            page.apply_redactions(images=pymupdf.PDF_REDACT_IMAGE_NONE, graphics=pymupdf.PDF_REDACT_LINE_ART_NONE, text=pymupdf.PDF_REDACT_TEXT_REMOVE)
            for cell, value in cell_results:
                r = cell["rect"]; rect = pymupdf.Rect(r.x0 + 1.2, r.y0 + 1.0, r.x1 - 1.2, r.y1 - 1.0)
                _insert_pdf_text_fitted(page, rect, value, PDFTextBlock(page_index, tuple(rect), cell["text"], 9.5, 0, 0))
    for block in blocks: source[block.page].add_redact_annot(pymupdf.Rect(block.bbox), fill=False, cross_out=False)
    for page in source: page.apply_redactions(images=pymupdf.PDF_REDACT_IMAGE_NONE, graphics=pymupdf.PDF_REDACT_LINE_ART_NONE, text=pymupdf.PDF_REDACT_TEXT_REMOVE)
    for block, value in zip(blocks, translated_blocks):
        rect = pymupdf.Rect(block.bbox); rect = pymupdf.Rect(rect.x0, rect.y0 - 0.5, rect.x1, rect.y1 + max(1.0, block.fontsize * 0.18))
        max_bottom = _pdf_block_available_bottom(block, blocks, source[block.page])
        _insert_pdf_text_fitted(source[block.page], rect, value, block, max_bottom=max_bottom)
    output = source.tobytes(garbage=4, deflate=True, clean=True); source.close()
    final_validation = _validate_final_pdf(output, original_block_texts + original_table_texts)
    if not final_validation["passed"]:
        raise RuntimeError("La reconstrucción PDF fue bloqueada: la validación final detectó contenido protegido que no sobrevivió al renderizado.")
    counts = _aggregate_counts(original_block_texts + original_table_texts)
    if table_infos: counts["protectedByType"]["table"] = len(table_infos)
    else: counts["protectedByType"].pop("table", None)
    return output, {"format": "pdf", "pages": page_count, "textBlocks": len(blocks), "tables": len(table_infos), "tableDetails": table_infos, "counts": counts, "validation": final_validation}

def _docx_media_inventory(path_or_bytes):
    if hasattr(path_or_bytes, "read"):
        raw = path_or_bytes.read()
        try: path_or_bytes.seek(0)
        except Exception: pass
        z = zipfile.ZipFile(BytesIO(raw))
    elif isinstance(path_or_bytes, (bytes, bytearray)):
        z = zipfile.ZipFile(BytesIO(path_or_bytes))
    else:
        z = zipfile.ZipFile(str(path_or_bytes))
    try:
        return {name for name in z.namelist() if name.startswith("word/media/") and not name.endswith("/")}
    finally:
        z.close()

def _set_paragraph_text_preserve_drawings(paragraph, value):
    from docx.oxml.ns import qn
    text_nodes = []
    for run in paragraph.runs:
        for child in run._r.iterchildren():
            if child.tag == qn("w:t"):
                text_nodes.append(child)
    if text_nodes:
        text_nodes[0].text = value
        for node in text_nodes[1:]: node.text = ""
    else:
        if value and not any(el.tag == qn("w:drawing") or el.tag == qn("w:pict") for el in paragraph._p.iter()):
            paragraph.add_run(value)

def _translate_docx_paragraphs(document, source_lang, target_lang, translate_fn):
    texts = []; locations = []
    def add_paragraph(paragraph):
        text = paragraph.text.strip()
        if text:
            texts.append(text); locations.append(paragraph)
    for paragraph in document.paragraphs: add_paragraph(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs: add_paragraph(paragraph)
    translated = []
    for start in range(0, len(texts), 18):
        batch = texts[start:start + 18]
        segments = [(f"{start + i + 1:06d}", str(text)) for i, text in enumerate(batch)]
        translated.extend(_translate_segments(segments, source_lang, target_lang, translate_fn))
    for paragraph, value in zip(locations, translated):
        _set_paragraph_text_preserve_drawings(paragraph, value)
    return texts, translated

def translate_docx_document(path, source_lang, target_lang, translate_fn):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Falta python-docx. La dependencia está declarada en requirements.txt.") from exc
    original_media = _docx_media_inventory(path)
    document = Document(str(path))
    original_texts, _ = _translate_docx_paragraphs(document, source_lang, target_lang, translate_fn)
    counts = _aggregate_counts(original_texts)
    output = BytesIO()
    document.save(output)
    data = output.getvalue()
    output_media = _docx_media_inventory(data)
    if original_media != output_media:
        missing = sorted(original_media - output_media)
        raise RuntimeError(f"La reconstrucción DOCX fue bloqueada: se perderían {len(missing)} imagen(es) del documento original.")
    image_count = len(original_media)
    return data, {"format": "docx", "pages": None, "textBlocks": len(original_texts), "tables": len(document.tables),
                  "images": image_count, "imagesPreserved": original_media == output_media,
                  "tableDetails": [{"number": i + 1, "page": None, "rows": len(table.rows), "columns": len(table.columns),
                                    "cells": len(table.rows) * len(table.columns), "size": f"{len(table.rows)}×{len(table.columns)}"}
                                   for i, table in enumerate(document.tables)],
                  "counts": counts, "validation": {"passed": True, "checked": image_count, "issues": []}}

def translate_document(path, source_lang, target_lang, translate_fn):
    suffix = Path(path).suffix.lower()
    if suffix == ".pdf":
        data, info = translate_pdf_document(path, source_lang, target_lang, translate_fn); return data, "pdf", info
    if suffix == ".docx":
        data, info = translate_docx_document(path, source_lang, target_lang, translate_fn); return data, "docx", info
    raise ValueError("Formato de documento no soportado.")